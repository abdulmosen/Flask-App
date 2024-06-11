from flask import Flask, request, redirect, url_for, render_template, send_from_directory
from werkzeug.utils import secure_filename
import os
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def docx_to_excel(docx_path, excel_path):
    doc = Document(docx_path)
    wb = Workbook()
    ws = wb.active
    ws.delete_rows(1, ws.max_row)
    row = 1
    for para in doc.paragraphs:
        if para.text.strip():
            ws.cell(row=row, column=1, value=para.text)
            ws.cell(row=row, column=1).alignment = Alignment(wrap_text=True)
            row += 1
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows, start=row):
            for col_idx, cell in enumerate(row.cells, start=1):
                ws.cell(row=row_idx, column=col_idx, value=cell.text)
                ws.cell(row=row_idx, column=col_idx).alignment = Alignment(wrap_text=True)
    for column_cells in ws.columns:
        max_length = 0
        column = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width
    wb.save(excel_path)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            docx_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(docx_path)
            excel_path = os.path.splitext(docx_path)[0] + '.xlsx'
            docx_to_excel(docx_path, excel_path)
            return redirect(url_for('uploaded_file', filename=os.path.basename(excel_path)))
    return render_template('index.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)
